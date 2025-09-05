import io
import os
from typing import Optional, Dict, Any
try:
    import PyPDF2
except ImportError:
    PyPDF2 = None

try:
    import docx
except ImportError:
    docx = None

try:
    import openpyxl
except ImportError:
    openpyxl = None
from io import BytesIO

class DocumentProcessor:
    def __init__(self):
        """Initialize the document processor"""
        self.supported_formats = ['.pdf', '.docx', '.xlsx', '.txt']
    
    def process_document(self, content: bytes, file_extension: str) -> Optional[str]:
        """Process a document and extract text content"""
        try:
            if file_extension.lower() == '.pdf':
                if PyPDF2 is not None:
                    return self._process_pdf(content)
                else:
                    print("PyPDF2 not available for PDF processing")
                    return None
            elif file_extension.lower() == '.docx':
                if docx is not None:
                    return self._process_docx(content)
                else:
                    print("python-docx not available for DOCX processing")
                    return None
            elif file_extension.lower() == '.xlsx':
                if openpyxl is not None:
                    return self._process_xlsx(content)
                else:
                    print("openpyxl not available for XLSX processing")
                    return None
            elif file_extension.lower() == '.txt':
                return self._process_txt(content)
            else:
                print(f"Unsupported file format: {file_extension}")
                return None
                
        except Exception as e:
            print(f"Error processing document: {str(e)}")
            return None
    
    def _process_pdf(self, content: bytes) -> Optional[str]:
        """Extract text from PDF content"""
        try:
            pdf_file = BytesIO(content)
            pdf_reader = PyPDF2.PdfReader(pdf_file)
            
            text_content = []
            for page_num in range(len(pdf_reader.pages)):
                page = pdf_reader.pages[page_num]
                page_text = page.extract_text()
                if page_text:
                    text_content.append(page_text)
            
            return '\n'.join(text_content)
            
        except Exception as e:
            print(f"Error processing PDF: {str(e)}")
            return None
    
    def _process_docx(self, content: bytes) -> Optional[str]:
        """Extract text from DOCX content"""
        try:
            docx_file = BytesIO(content)
            doc = docx.Document(docx_file)
            
            text_content = []
            
            # Extract text from paragraphs
            for paragraph in doc.paragraphs:
                if paragraph.text.strip():
                    text_content.append(paragraph.text)
            
            # Extract text from tables
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        text_content.append(' | '.join(row_text))
            
            return '\n'.join(text_content)
            
        except Exception as e:
            print(f"Error processing DOCX: {str(e)}")
            return None
    
    def _process_xlsx(self, content: bytes) -> Optional[str]:
        """Extract text from XLSX content"""
        try:
            xlsx_file = BytesIO(content)
            workbook = openpyxl.load_workbook(xlsx_file, read_only=True)
            
            text_content = []
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                text_content.append(f"Sheet: {sheet_name}")
                
                # Extract data from cells
                for row in sheet.iter_rows(values_only=True):
                    row_values = []
                    for cell_value in row:
                        if cell_value is not None:
                            row_values.append(str(cell_value))
                    
                    if row_values:
                        text_content.append(' | '.join(row_values))
                
                text_content.append("")  # Add blank line between sheets
            
            workbook.close()
            return '\n'.join(text_content)
            
        except Exception as e:
            print(f"Error processing XLSX: {str(e)}")
            return None
    
    def _process_txt(self, content: bytes) -> Optional[str]:
        """Extract text from TXT content"""
        try:
            # Try different encodings
            encodings = ['utf-8', 'utf-16', 'latin-1', 'cp1252']
            
            for encoding in encodings:
                try:
                    return content.decode(encoding)
                except UnicodeDecodeError:
                    continue
            
            # If all encodings fail, decode with errors='ignore'
            return content.decode('utf-8', errors='ignore')
            
        except Exception as e:
            print(f"Error processing TXT: {str(e)}")
            return None
    
    def extract_metadata(self, content: bytes, file_extension: str) -> Dict[str, Any]:
        """Extract metadata from documents"""
        metadata = {
            'file_type': file_extension,
            'size_bytes': len(content)
        }
        
        try:
            if file_extension.lower() == '.pdf':
                pdf_file = BytesIO(content)
                pdf_reader = PyPDF2.PdfReader(pdf_file)
                
                metadata.update({
                    'pages': len(pdf_reader.pages),
                    'title': pdf_reader.metadata.get('/Title', '') if pdf_reader.metadata else '',
                    'author': pdf_reader.metadata.get('/Author', '') if pdf_reader.metadata else '',
                    'subject': pdf_reader.metadata.get('/Subject', '') if pdf_reader.metadata else ''
                })
                
            elif file_extension.lower() == '.docx':
                docx_file = BytesIO(content)
                doc = docx.Document(docx_file)
                
                metadata.update({
                    'paragraphs': len(doc.paragraphs),
                    'tables': len(doc.tables),
                    'title': doc.core_properties.title or '',
                    'author': doc.core_properties.author or '',
                    'subject': doc.core_properties.subject or ''
                })
                
            elif file_extension.lower() == '.xlsx':
                xlsx_file = BytesIO(content)
                workbook = openpyxl.load_workbook(xlsx_file, read_only=True)
                
                metadata.update({
                    'sheets': len(workbook.sheetnames),
                    'sheet_names': workbook.sheetnames
                })
                
                workbook.close()
                
        except Exception as e:
            print(f"Error extracting metadata: {str(e)}")
        
        return metadata
    
    def is_supported_format(self, file_extension: str) -> bool:
        """Check if the file format is supported"""
        return file_extension.lower() in self.supported_formats
    
    def get_supported_formats(self) -> list:
        """Get list of supported file formats"""
        return self.supported_formats.copy()
